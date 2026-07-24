"""Generate a client-ready Word document from the deliverables status."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x1F, 0x4D, 0x3F)
GOLD = RGBColor(0xA8, 0x79, 0x2C)
INK = RGBColor(0x1A, 0x21, 0x1E)
SOFT = RGBColor(0x4A, 0x55, 0x4F)
GOOD = RGBColor(0x2E, 0x73, 0x58)
WARN = RGBColor(0x9A, 0x6B, 0x14)
IDLE = RGBColor(0x8A, 0x94, 0x90)

doc = Document()

# base style
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(10.5)
st.font.color.rgb = INK

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def spacer(pts=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p

def heading(text, size=15, color=GREEN, before=16, after=4, rule=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    if rule:
        pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'4')
        bottom.set(qn('w:color'),'1F4D3F'); pbdr.append(bottom); pPr.append(pbdr)
    return p

def para(text, color=SOFT, size=10.5, bold=False, after=6, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

STATUS = {'done': ('Done', GOOD), 'partial': ('Partial', WARN), 'pending': ('Pending', IDLE)}

def status_table(rows, cols):
    """rows: list of tuples matching cols. cols: list of header names.
    A cell value of ('S','done') renders a status."""
    t = doc.add_table(rows=1, cols=len(cols)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'; t.autofit = True
    hdr = t.rows[0].cells
    for i, c in enumerate(cols):
        shade(hdr[i], '1F4D3F')
        run = hdr[i].paragraphs[0].add_run(c); run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]; p = cell.paragraphs[0]
            if isinstance(val, tuple) and val[0] == 'S':
                label, col = STATUS[val[1]]
                run = p.add_run('● ' + label); run.font.color.rgb = col
                run.bold = True; run.font.size = Pt(9)
            else:
                run = p.add_run(str(val)); run.font.size = Pt(9)
                if i == 0 and cols[0] == 'ID':
                    run.bold = True; run.font.color.rgb = INK
                else:
                    run.font.color.rgb = SOFT
    return t

# ---------------- TITLE ----------------
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
r = p.add_run('DEVELOPMENT STATUS REPORT'); r.bold = True; r.font.size = Pt(9)
r.font.color.rgb = GOLD
# letter spacing
rPr = r._element.get_or_add_rPr(); spc = OxmlElement('w:spacing'); spc.set(qn('w:val'),'40'); rPr.append(spc)

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
r = p.add_run('BHITR CRM — Deliverables & Build Status'); r.bold = True
r.font.size = Pt(22); r.font.color.rgb = GREEN

para('A traceable, section-by-section account of what is built and working today, and what '
     'remains — measured against the Product Requirements Document v1.0.', size=11, after=10)

# meta line
t = doc.add_table(rows=1, cols=3)
meta = [('Source of truth', 'PRD v1.0 (02 Jul 2026)'),
        ('Platform', 'Django + PostgreSQL, deployed'),
        ('Updated', '11 July 2026')]
for i,(k,v) in enumerate(meta):
    c = t.rows[0].cells[i]
    rp = c.paragraphs[0].add_run(k.upper()); rp.font.size = Pt(7.5); rp.font.color.rgb = IDLE; rp.bold = True
    c.add_paragraph().add_run(v).font.size = Pt(9.5)
# remove borders on meta table
tbl = t._tbl
for el in tbl.iter():
    if el.tag == qn('w:tcBorders'):
        el.getparent().remove(el)
spacer(4)

# legends
heading('How to read this report', 12, GREEN, before=12, after=6)
para('Status:  ● Done — built & working, verified in code   |   '
     '● Partial — basic version live, PRD depth pending   |   '
     '● Pending — not started', size=9.5, after=3)
para('Dependency:  Development — fully in our control   |   '
     'Vendor — needs a 3rd-party paid service/account   |   '
     'Sign-off — a business/legal decision, not code', size=9.5, after=8)

# callout
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
p.paragraph_format.left_indent = Inches(0.1)
r = p.add_run('Is everything buildable?  '); r.bold = True; r.font.color.rgb = GREEN; r.font.size = Pt(10)
r = p.add_run('Yes — nothing in the PRD is architecturally impossible on our stack. Items marked '
              '“Vendor” need the client to choose & pay for an external service (WhatsApp, telephony, '
              'e-sign, screening, OCR, accounting); items marked “Sign-off” need a business or legal '
              'decision (hosting region, compliance values, penetration test), not development.')
r.font.color.rgb = SOFT; r.font.size = Pt(10)

# ---------------- SUMMARY ----------------
heading('Summary at a glance', 15, GREEN, rule=True)
status_table([
    ('Phase 1 — Production MVP', ('S','partial'), 'Core CRUD, 5 roles, dashboards, leads, documents & tasks live. Role-model depth, Client/Case/Bank-Application model, SLA engine, compliance gate & audit infrastructure pending. (~40%)'),
    ('Phase 2 — Ops depth, invoicing, partners, HR', ('S','pending'), 'Ops subflows, invoicing & commission, channel partners, automation, HR, WhatsApp/telephony.'),
    ('Phase 3 — Intelligence & compliance automation', ('S','pending'), 'Screening, OCR, e-sign, report builder, accounting export.'),
    ('Phase 4 — Portals & expansion', ('S','pending'), 'Client & partner portals, mobile app, industry packs.'),
], ['Phase', 'Status', 'Detail'])
spacer(4)
para('Delivered beyond the PRD (client-requested): CEO Customization revenue sheet · per-lead Audit '
     'Log · custom date-range filter · CEO-only delete · deployment pipeline.', color=GREEN, size=9.5, bold=True)

# ---------------- PHASE 1 ----------------
heading('Phase 1 — Production MVP  (in progress, ~40%)', 15, GREEN, rule=True)

heading('1. Authentication & Access', 12, INK)
status_table([
    ('FR-01','Login, TOTP 2FA, lockout, session timeout, device history',('S','partial'),'Development'),
    ('AD-01','User admin: invite, deactivate, force reset, kill sessions, IP allowlist',('S','partial'),'Development'),
    ('AD-02','Role editor + field-level security',('S','partial'),'Development'),
], ['ID','Requirement','Status','Dep'])

heading('1b. User Roles & Permissions (Section 8)', 12, INK)
para('The PRD ships 13 default roles with Own/Team/All hierarchy, field-level security, delegation '
     'and segregation of duties. The current app has 5 roles and a module-level matrix.', size=9.5, after=4)
status_table([
    ('Role catalogue: 13 roles',('S','partial'),'Have 5 (CEO, Sales Director, Ops Manager, Advisor, Accountant). Missing 8: Super Admin, Team Leader, Operations Executive, HR Executive, Compliance Officer, Marketing Executive, Telecaller, External Auditor'),
    ('Module permission matrix',('S','partial'),'Editable & persisted; ~300 per-action switches pending'),
    ('Field-level security',('S','pending'),'Hide bank commission % from advisors; salary HR+CEO only; masked identity docs'),
    ('Own / Team / All visibility hierarchy',('S','partial'),'Own & All work; Team scope pending'),
    ('Overrides · delegation · segregation of duties · watchers',('S','pending'),'Single role only; approval-workflow dependent'),
], ['Requirement','Status','Notes'])

heading('2. Lead Management', 12, INK)
status_table([
    ('LM-01','Seven-step lead wizard, extended fields',('S','partial'),'Development'),
    ('LM-02','Web-to-lead API endpoint + spam controls',('S','pending'),'Development'),
    ('LM-05','Instant eligibility check (DBR, LTV cap, cash-to-close)',('S','pending'),'Dev + Sign-off'),
    ('LM-06','Duplicate block + fuzzy warning on save',('S','pending'),'Development'),
    ('LM-08','Assignment rules engine (round-robin, caps)',('S','pending'),'Development'),
    ('LM-09','First-contact SLA countdown + escalation',('S','pending'),'Development'),
    ('LM-15','CSV import wizard (mapping, dedupe, undo)',('S','pending'),'Development'),
    ('LM-17','Lost reasons + Pareto; reopen with approval',('S','partial'),'Development'),
    ('LM-18','Consent capture per channel',('S','pending'),'Development'),
    ('—','List, edit, delete, bulk actions, export, pipeline, sources, notes',('S','done'),'Development'),
], ['ID','Requirement','Status','Dep'])

heading('3. Pipeline & Cases', 12, INK)
status_table([
    ('PL-01','Lead → Client + Case conversion (foundational — no separate entity yet)',('S','pending'),'Development'),
    ('PL-02','Configurable stages + entry/exit gates + target days',('S','partial'),'Development'),
    ('PL-04','In-stage aging bands + rotting badges',('S','done'),'Development'),
    ('PL-05','On Hold / Declined / Withdrawn + SLA pause',('S','partial'),'Development'),
    ('PL-09','Stage history timestamps for TAT analytics',('S','partial'),'Development'),
    ('PL-10','Case reference numbering',('S','pending'),'Development'),
], ['ID','Requirement','Status','Dep'])

heading('4. Bank Applications', 12, INK)
status_table([
    ('BA-01–06','Multiple bank applications per case — independent status, refs, follow-up dates, rejection Pareto, timestamps (needs Case model; currently one lead = one bank)',('S','pending'),'Development'),
], ['ID','Requirement','Status','Dep'])

heading('5. Operations Processing', 12, INK)
status_table([
    ('OPS-01','Handover gate (documents, KYC, data completeness)',('S','pending'),'Development'),
    ('OPS-03','Per-document verify/reject loop + auto chase tasks',('S','partial'),'Development'),
    ('OPS-05','Follow-up log + 3/7-day silence rules',('S','pending'),'Development'),
    ('OPS-07','Pre-approval capture + validity countdown',('S','pending'),'Development'),
    ('OPS-09','FOL structured terms + fixed-period end date',('S','pending'),'Development'),
    ('OPS-13','Deed / disbursement capture → triggers finance',('S','pending'),'Development'),
    ('OPS-08,10-12','Valuation / Buyout / NOC / Transfer subflows (Phase 2)',('S','pending'),'Development'),
], ['ID','Requirement','Status','Dep'])

heading('6–8. Client Lifecycle · Tasks · Documents', 12, INK)
status_table([
    ('CL-01','Client entity + 360 view (no Client entity yet)',('S','pending'),'Development'),
    ('TA-01','Task object: types, links, reminders, outcomes',('S','partial'),'Development'),
    ('TA-08','Unified activity timeline + @mentions on all records',('S','partial'),'Development'),
    ('—','Task list, overdue, create, complete, export',('S','done'),'Development'),
    ('DM-03','Document statuses + verified-by (verified-at pending)',('S','partial'),'Development'),
    ('DM-06','Expiry engine (30/14/7-day alerts)',('S','pending'),'Development'),
    ('DM-11','Virus scan, size/type limits',('S','pending'),'Vendor'),
    ('—','Upload (lead form + per-lead + title deed), verify/reject/reupload, export',('S','done'),'Development'),
], ['ID','Requirement','Status','Dep'])

heading('9. Compliance & Audit', 12, INK)
status_table([
    ('CO-01','KYC checklist + Compliance-only gate blocking submission',('S','pending'),'Development'),
    ('CO-05','Consent objects per channel + do-not-contact flag',('S','pending'),'Development'),
    ('CO-09','Append-only audit + recycle bin (lead-level audit live; system-wide, hash-chain, recycle bin & no-hard-delete pending)',('S','partial'),'Development'),
], ['ID','Requirement','Status','Dep'])

heading('10–12. Notifications · Reporting · Settings', 12, INK)
status_table([
    ('NA-01','In-app notification center + preferences (bell decorative today)',('S','pending'),'Development'),
    ('NA-02','Email + push delivery; digest + quiet hours',('S','pending'),'Vendor'),
    ('NA-07','SLA engine (business-hours + holidays + hold-pause)',('S','pending'),'Development'),
    ('RP-01','Role dashboards on live data',('S','partial'),'Development'),
    ('RP-03','Report catalogue (top 8 in Phase 1)',('S','partial'),'Development'),
    ('AD-03','Masters admin (banks, doc types, sources) + safe deactivate',('S','partial'),'Development'),
    ('AD-06','Company profile, branding, numbering (only personal profile saves today)',('S','pending'),'Development'),
    ('IN-01/02','REST API + web-to-lead',('S','pending'),'Development'),
], ['ID','Requirement','Status','Dep'])

# ---------------- PHASE 2-4 ----------------
doc.add_page_break()
heading('Phases 2–4 — Planned', 15, GREEN, rule=True)
heading('Phase 2 — Operational depth & engagement', 12, INK)
para('Ops subflows (Valuation, Buyout, NOC, Transfer) · Channel Partners end-to-end · Invoicing → '
     'receivables → payout runs → incentives · No-code automation & approvals · HR (attendance, leave, '
     'targets) · WhatsApp, telephony, mailbox/calendar (Vendor) · Milestone & post-close messaging · '
     'Template studio, custom fields, sandbox.', size=10)
heading('Phase 3 — Intelligence & compliance automation', 12, INK)
para('Screening API, OCR, E-signature (Vendor) · Risk rating / EDD / UBO / DSR / retention · Lead scoring '
     'v2 + weighted forecast · Report builder & scheduling · Accounting export (Zoho / QuickBooks / Tally) + '
     'payment links (Vendor).', size=10)
heading('Phase 4 — Portals & expansion', 12, INK)
para('Client portal + mobile app · Partner portal · Marketing campaign module · Industry packs '
     '(real-estate / insurance) · Multi-company groundwork.', size=10)

# ---------------- PRIORITY ORDER ----------------
heading('Recommended order to complete Phase 1', 15, GREEN, rule=True)
order = [
    'Full role model — 8 missing roles, field-level security, Own/Team/All hierarchy, per-user overrides',
    'Client + Case + Bank Application data model — foundational for all operations & finance',
    'No-hard-delete + recycle bin + system-wide audit — a PRD non-negotiable',
    '2FA + session/login history + account lockout',
    'SLA engine + first-contact SLA + “every open lead has a future task”',
    'Operations workflow — handover gate, follow-up log, pre-approval / FOL / disbursement capture',
    'KYC hard gate + consent capture',
    'Eligibility check, deduplication, assignment rules, import wizard',
    'Notification center + remaining role dashboards',
    'REST API + web-to-lead; save all lead-wizard fields; company/branding settings',
]
for i, item in enumerate(order, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{i}.  '); r.bold = True; r.font.color.rgb = GOLD; r.font.size = Pt(10)
    r = p.add_run(item); r.font.color.rgb = SOFT; r.font.size = Pt(10)

# ---------------- DEPENDENCIES ----------------
heading('External dependencies the client must provide', 15, GREEN, rule=True)
para('These features will be built by us, but require a third-party account/subscription (Vendor) or a '
     'business/legal decision (Sign-off) to operate.', size=10, after=6)
status_table([
    ('WhatsApp Business API (BSP)','Vendor','WhatsApp notifications, document requests, shared inbox'),
    ('Telephony / CPaaS','Vendor','Click-to-call, screen-pop, call recording'),
    ('Transactional email & SMS','Vendor','Email/SMS notifications & digests'),
    ('E-signature','Vendor','Agreements, FOL signing'),
    ('Screening data provider','Vendor','Sanctions / PEP checks'),
    ('OCR service','Vendor','Auto-read passport / EID / salary certificate'),
    ('Accounting connector','Vendor','Zoho Books / QuickBooks / Tally export'),
    ('Google / Microsoft workspace','Vendor','SSO, mailbox & calendar sync'),
    ('Google Maps','Vendor','Attendance geofence, address normalization'),
    ('UAE hosting region','Sign-off','PDPL residency — verify region or move (e.g. AWS Bahrain)'),
    ('Compliance values','Sign-off','LTV caps, DBR, retention years — verified by compliance advisor'),
    ('Penetration test','Sign-off','Security audit before go-live & annually'),
], ['Need','Type','For'])

spacer(10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BHITR CRM · Development Status Report · Prepared 11 July 2026 · Confidential')
r.font.size = Pt(8); r.font.color.rgb = IDLE

doc.save('BHITR_CRM_Deliverables_Status.docx')
print('saved BHITR_CRM_Deliverables_Status.docx')
